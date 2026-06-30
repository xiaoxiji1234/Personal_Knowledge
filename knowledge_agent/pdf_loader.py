from __future__ import annotations

import csv
from io import StringIO
from pathlib import Path


def load_document_text(path: Path) -> tuple[str, dict[str, object]]:
    """Load supported document formats into plain text plus parser metadata."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix == ".xlsx":
        return _load_xlsx(path)
    if suffix == ".xls":
        return _load_xls(path)
    if suffix == ".csv":
        return _load_csv(path)
    raw = path.read_bytes()
    return _decode_bytes(raw), {"parser": "text", "pages": 1, "quality": "text"}


def _load_pdf(path: Path) -> tuple[str, dict[str, object]]:
    """Extract text from PDF pages using pypdf, with a bytes fallback if missing."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raw = path.read_bytes()
        return _decode_bytes(raw), {"parser": "bytes-fallback", "pages": 0, "quality": "unknown"}

    reader = PdfReader(str(path))
    pages: list[str] = []
    empty_pages = 0
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if not page_text.strip():
            empty_pages += 1
        pages.append(page_text)
    quality = "scanned-or-empty" if pages and empty_pages == len(pages) else "text"
    return "\n\n".join(pages), {"parser": "pypdf", "pages": len(reader.pages), "quality": quality}


def _load_docx(path: Path) -> tuple[str, dict[str, object]]:
    """Extract paragraphs and table cells from a Word .docx file."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to parse .docx files.") from exc

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    quality = "text" if parts else "empty"
    return "\n\n".join(parts), {"parser": "python-docx", "pages": 1, "quality": quality}


def _load_xlsx(path: Path) -> tuple[str, dict[str, object]]:
    """Extract visible cell values from an Excel .xlsx workbook."""
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required to parse .xlsx files.") from exc

    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [_cell_to_text(value) for value in row]
            values = [value for value in values if value]
            if values:
                parts.append(" | ".join(values))
    workbook.close()
    non_header_rows = [part for part in parts if not part.startswith("# Sheet:")]
    quality = "text" if non_header_rows else "empty"
    return "\n".join(parts), {"parser": "openpyxl", "pages": len(workbook.worksheets), "quality": quality}


def _load_csv(path: Path) -> tuple[str, dict[str, object]]:
    """Extract rows from a CSV file into pipe-delimited text."""
    text = _decode_bytes(path.read_bytes())
    rows = csv.reader(StringIO(text))
    parts = [" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows]
    parts = [part for part in parts if part]
    quality = "text" if parts else "empty"
    return "\n".join(parts), {"parser": "csv", "pages": 1, "quality": quality}


def _load_xls(path: Path) -> tuple[str, dict[str, object]]:
    """Extract cell values from a legacy Excel .xls workbook."""
    try:
        import xlrd
    except ImportError as exc:
        raise RuntimeError("xlrd is required to parse .xls files.") from exc

    workbook = xlrd.open_workbook(str(path))
    parts: list[str] = []
    for sheet in workbook.sheets():
        parts.append(f"# Sheet: {sheet.name}")
        for row_index in range(sheet.nrows):
            values = [_cell_to_text(sheet.cell_value(row_index, column_index)) for column_index in range(sheet.ncols)]
            values = [value for value in values if value]
            if values:
                parts.append(" | ".join(values))
    non_header_rows = [part for part in parts if not part.startswith("# Sheet:")]
    quality = "text" if non_header_rows else "empty"
    return "\n".join(parts), {"parser": "xlrd", "pages": workbook.nsheets, "quality": quality}


def _cell_to_text(value: object) -> str:
    """Convert spreadsheet cell values to stable plain text."""
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _decode_bytes(raw: bytes) -> str:
    """Decode text-like bytes with common encodings used by uploaded files."""
    for encoding in ("utf-8", "utf-16", "gb18030", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")
